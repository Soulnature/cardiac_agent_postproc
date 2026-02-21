import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Simple markdown parser
    lines = content.split('\n')
    current_list_level = 0
    in_code_block = False

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        # List items
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Ordered list
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # Bold text handling (simple version) inside paragraphs
        else:
            # Handle mixed bold text
            bold_parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in bold_parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Created {docx_path}")

if __name__ == "__main__":
    markdown_to_docx(
        "/data484_5/xzhao14/cardiac_agent_postproc/miccai_method_section.md",
        "/data484_5/xzhao14/cardiac_agent_postproc/miccai_method_section.docx"
    )
